"""
📄 OMNIMIND - Advanced Document & OCR Processor

Handles PDF, DOCX, images, and OCR with superior extraction capabilities.
"""

import os
import io
from typing import Dict, Any, List, Optional
from pathlib import Path
import asyncio

# Document processing
PDFPLUMBER_AVAILABLE = False
PANDAS_AVAILABLE = False

try:
    import PyPDF2
    from docx import Document
    import openpyxl
    print("[INFO] Document processing libraries loaded")
except ImportError as e:
    print(f"[WARNING] Some document libraries not available: {e}")

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    print("[INFO] pdfplumber not available - using PyPDF2 only")

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    print("[INFO] pandas not available - Excel/CSV features limited")

# OCR and image processing
PYTESSERACT_AVAILABLE = False
try:
    from PIL import Image
    print("[INFO] PIL/Pillow loaded")
except ImportError:
    print("[WARNING] PIL/Pillow not available")

try:
    import pytesseract
    PYTESSERACT_AVAILABLE = True
except ImportError:
    print("[INFO] pytesseract not available - OCR features limited")


class DocumentProcessor:
    """Advanced document processing with OCR capabilities"""
    
    def __init__(self):
        self.max_file_size = int(os.getenv("MAX_FILE_SIZE_MB", "10")) * 1024 * 1024
        self.enable_ocr = os.getenv("ENABLE_OCR", "true").lower() == "true"
        
        # Configure Tesseract (for OCR)
        if self.enable_ocr:
            self._configure_tesseract()
    
    def _configure_tesseract(self):
        """Configure Tesseract OCR"""
        # Common Tesseract paths on Windows
        possible_paths = [
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
            r"C:\Users\Public\Tesseract-OCR\tesseract.exe"
        ]
        
        for path in possible_paths:
            if os.path.exists(path):
                pytesseract.pytesseract.tesseract_cmd = path
                print(f"✅ Tesseract OCR configured: {path}")
                return
        
        print("[WARNING] Tesseract not found. OCR features will be limited.")
        print("   Install from: https://github.com/UB-Mannheim/tesseract/wiki")
    
    async def process_file(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """
        Process any supported file type and extract content.
        
        Args:
            file_path: Path to the file
            file_type: MIME type or extension
        
        Returns:
            Dictionary with extracted content and metadata
        """
        
        # Check file size
        file_size = os.path.getsize(file_path)
        if file_size > self.max_file_size:
            return {
                "success": False,
                "error": f"File too large. Max size: {self.max_file_size / 1024 / 1024}MB"
            }
        
        # Route to appropriate processor
        if file_type in ["application/pdf", ".pdf"]:
            return await self._process_pdf(file_path)
        elif file_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", ".docx"]:
            return await self._process_docx(file_path)
        elif file_type in ["text/plain", ".txt", ".md"]:
            return await self._process_text(file_path)
        elif file_type in ["application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", ".xlsx"]:
            return await self._process_excel(file_path)
        elif file_type in ["text/csv", ".csv"]:
            return await self._process_csv(file_path)
        elif file_type in ["image/png", "image/jpeg", "image/jpg", ".png", ".jpg", ".jpeg"]:
            return await self._process_image(file_path)
        else:
            return {
                "success": False,
                "error": f"Unsupported file type: {file_type}"
            }
    
    async def _process_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract text and tables from PDF"""
        try:
            text_content = []
            tables = []
            
            # Try pdfplumber first (Better for tables)
            if PDFPLUMBER_AVAILABLE:
                with pdfplumber.open(file_path) as pdf:
                    for page_num, page in enumerate(pdf.pages, 1):
                        page_text = page.extract_text()
                        if page_text:
                            # Clean garbage text (null bytes, excessive control chars)
                            clean_text = "".join(ch for ch in page_text if ch.isprintable() or ch in '\n\t')
                            if len(clean_text) > 10: # Min content check
                                text_content.append(f"--- Page {page_num} ---\n{clean_text}")
                        
                        page_tables = page.extract_tables()
                        if page_tables:
                            for table in page_tables:
                                tables.append({"page": page_num, "data": table})
            
            # Fallback to PyPDF2
            else:
                print("[INFO] Using PyPDF2 fallback")
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    for page_num, page in enumerate(reader.pages, 1):
                        page_text = page.extract_text()
                        if page_text:
                            clean_text = "".join(ch for ch in page_text if ch.isprintable() or ch in '\n\t')
                            if len(clean_text) > 10:
                                text_content.append(f"--- Page {page_num} ---\n{clean_text}")
            
            return {
                "success": True,
                "type": "pdf",
                "text": "\n\n".join(text_content),
                "tables": tables,
                "pages": len(text_content),
                "summary": f"Extracted {len(text_content)} pages and {len(tables)} tables"
            }
        
        except Exception as e:
            return {
                "success": False,
                "error": f"PDF processing error: {str(e)}"
            }
    
    async def _process_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract text from DOCX"""
        try:
            doc = Document(file_path)
            
            # Extract paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            text_content = "\n\n".join(paragraphs)
            
            return {
                "success": True,
                "type": "docx",
                "text": text_content,
                "tables": tables,
                "paragraphs": len(paragraphs),
                "summary": f"Extracted {len(paragraphs)} paragraphs and {len(tables)} tables"
            }
        
        except Exception as e:
            return {
                "success": False,
                "error": f"DOCX processing error: {str(e)}"
            }
    
    async def _process_text(self, file_path: str) -> Dict[str, Any]:
        """Read plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            lines = content.split('\n')
            
            return {
                "success": True,
                "type": "text",
                "text": content,
                "lines": len(lines),
                "summary": f"Read {len(lines)} lines"
            }
        
        except Exception as e:
            return {
                "success": False,
                "error": f"Text file error: {str(e)}"
            }
    
    async def _process_excel(self, file_path: str) -> Dict[str, Any]:
        """Extract data from Excel"""
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            sheets_data = {}
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                sheets_data[sheet_name] = {
                    "rows": len(df),
                    "columns": list(df.columns),
                    "preview": df.head(10).to_dict('records')
                }
            
            # Create text summary
            text_summary = []
            for sheet_name, data in sheets_data.items():
                text_summary.append(f"Sheet: {sheet_name}")
                text_summary.append(f"Rows: {data['rows']}, Columns: {', '.join(data['columns'])}")
                text_summary.append("")
            
            return {
                "success": True,
                "type": "excel",
                "text": "\n".join(text_summary),
                "sheets": sheets_data,
                "summary": f"Extracted {len(sheets_data)} sheets"
            }
        
        except Exception as e:
            return {
                "success": False,
                "error": f"Excel processing error: {str(e)}"
            }
    
    async def _process_csv(self, file_path: str) -> Dict[str, Any]:
        """Extract data from CSV"""
        try:
            df = pd.read_csv(file_path)
            
            text_summary = [
                f"Rows: {len(df)}",
                f"Columns: {', '.join(df.columns)}",
                "",
                "Preview:",
                df.head(10).to_string()
            ]
            
            return {
                "success": True,
                "type": "csv",
                "text": "\n".join(text_summary),
                "rows": len(df),
                "columns": list(df.columns),
                "preview": df.head(10).to_dict('records'),
                "summary": f"Extracted {len(df)} rows, {len(df.columns)} columns"
            }
        
        except Exception as e:
            return {
                "success": False,
                "error": f"CSV processing error: {str(e)}"
            }
    
    async def _process_image(self, file_path: str) -> Dict[str, Any]:
        """
        Process image with OCR to extract text.
        This is a POWERFUL feature not available in standard ChatGPT!
        """
        try:
            # Open image
            image = Image.open(file_path)
            
            # Get image info
            width, height = image.size
            format_type = image.format
            
            # Perform OCR if enabled
            extracted_text = ""
            if self.enable_ocr:
                try:
                    extracted_text = pytesseract.image_to_string(image)
                except Exception as ocr_error:
                    extracted_text = f"OCR Error: {str(ocr_error)}\nNote: Install Tesseract OCR for text extraction."
            
            return {
                "success": True,
                "type": "image",
                "text": extracted_text,
                "width": width,
                "height": height,
                "format": format_type,
                "summary": f"Image: {width}x{height} {format_type}. Extracted {len(extracted_text)} characters via OCR."
            }
        
        except Exception as e:
            return {
                "success": False,
                "error": f"Image processing error: {str(e)}"
            }
    
    async def process_multiple_files(self, file_paths: List[str]) -> Dict[str, Any]:
        """Process multiple files concurrently"""
        tasks = []
        
        for file_path in file_paths:
            # Detect file type
            ext = Path(file_path).suffix.lower()
            tasks.append(self.process_file(file_path, ext))
        
        results = await asyncio.gather(*tasks)
        
        # Combine all text content
        combined_text = []
        all_results = []
        
        for i, result in enumerate(results):
            if result.get("success"):
                combined_text.append(f"=== File {i+1}: {file_paths[i]} ===")
                combined_text.append(result.get("text", ""))
                combined_text.append("")
                all_results.append(result)
        
        return {
            "success": True,
            "combined_text": "\n".join(combined_text),
            "individual_results": all_results,
            "total_files": len(file_paths),
            "successful": len(all_results)
        }
